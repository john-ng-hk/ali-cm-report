import json
import os
import shutil
from datetime import datetime, timedelta
from aliyunsdkcore.client import AcsClient
from aliyunsdkcms.request.v20190101.DescribeMetricListRequest import DescribeMetricListRequest
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
import pytz
import csv

# Load environment variables

def load_credentials():
    """Load credentials from AccessKey.csv file"""
    try:
        with open('AccessKey.csv', 'r') as file:
            # Read the first line to check headers
            first_line = file.readline().strip()
            print(f"\nCSV Headers found: {first_line}")
            
            # Reset file pointer
            file.seek(0)
            
            # Try different possible header formats
            possible_headers = [
                ['AccessKeyId', 'AccessKeySecret'],
                ['access_key_id', 'access_key_secret'],
                ['Access Key ID', 'Access Key Secret'],
                ['access key id', 'access key secret']
            ]
            
            reader = csv.DictReader(file)
            credentials = next(reader)  # Get the first row
            
            # Debug print all columns found
            print("\nColumns found in CSV:")
            for col in reader.fieldnames:
                print(f"- {col}")
            
            # Try to find credentials using different possible column names
            access_key_id = None
            access_key_secret = None
            
            for header_pair in possible_headers:
                if header_pair[0] in credentials and header_pair[1] in credentials:
                    access_key_id = credentials[header_pair[0]]
                    access_key_secret = credentials[header_pair[1]]
                    print(f"\nFound credentials using headers: {header_pair}")
                    break
            
            if not access_key_id or not access_key_secret:
                # Try direct column access as fallback
                if len(credentials) >= 2:
                    access_key_id = list(credentials.values())[0]
                    access_key_secret = list(credentials.values())[1]
                    print("\nFound credentials using first two columns")
            
            # Debug print found credentials
            print("\nFound credentials:")
            print(f"AccessKeyId: {access_key_id if access_key_id else 'Not found'}")
            print(f"AccessKeySecret: {'*' * len(access_key_secret) if access_key_secret else 'Not found'}")
            
            if not access_key_id or not access_key_secret:
                raise ValueError("Missing credentials in CSV file. Please ensure both AccessKeyId and AccessKeySecret are present.")
            
            return {
                'ACCESS_KEY_ID': access_key_id,
                'ACCESS_KEY_SECRET': access_key_secret
            }
    except FileNotFoundError:
        raise FileNotFoundError("AccessKey.csv file not found. Please ensure it exists in the current directory.")
    except Exception as e:
        raise Exception(f"Error reading AccessKey.csv: {str(e)}")

# Load credentials from AccessKey.csv
try:
    credentials = load_credentials()
    ACCESS_KEY_ID = credentials['ACCESS_KEY_ID']
    ACCESS_KEY_SECRET = credentials['ACCESS_KEY_SECRET']
    REGION_ID = 'cn-hongkong'
    
    # Validate credentials before creating client
    if not ACCESS_KEY_ID or not ACCESS_KEY_SECRET:
        raise ValueError("Invalid credentials: AccessKeyId or AccessKeySecret is empty")
    
    print("\nCredentials loaded successfully!")
except Exception as e:
    print(f"\nError: {str(e)}")
    print("\nPlease ensure your AccessKey.csv file:")
    print("1. Exists in the current directory")
    print("2. Has the correct column headers: AccessKeyId,AccessKeySecret")
    print("3. Contains valid credentials")
    print("\nExample AccessKey.csv format:")
    print("AccessKeyId,AccessKeySecret")
    print("your_access_key_id,your_access_key_secret")
    exit(1)

# Charts directory configuration
CHARTS_DIR = 'charts'

# Sprint Configuration
SPRINT_15_START = datetime(2025, 2, 13)  # Sprint 15 starts on 13 Feb 2025
SPRINT_DURATION = timedelta(days=14)  # Each sprint is 14 days

# Instance Configuration
INSTANCES = {
    'DEV': {
        'web': [{'id': 'i-xxx', 'name': 'DEV-WEB'}],
        'app': [{'id': 'i-xxx', 'name': 'DEV-APP'}],
        'rds': [{'id': 'rm-xxx', 'name': 'DEV-RDS'}]
    },
    'UAT': {
        'web': [
            {'id': 'i-xxx', 'name': 'UAT-WEB-1'},
            {'id': 'i-xxx', 'name': 'UAT-WEB-2'}
        ],
        'app': [
            {'id': 'i-xxx', 'name': 'UAT-APP-1'},
            {'id': 'i-xxx', 'name': 'UAT-APP-2'}
        ],
        'rds': [{'id': 'rm-xxx', 'name': 'UAT-RDS'}]
    }
}

# Metric configurations
ECS_METRICS = {
    'cpu': {
        'namespace': 'acs_ecs_dashboard',
        'metric_name': 'CPUUtilization',
        'unit': '%'
    },
    'memory': {
        'namespace': 'acs_ecs_dashboard',
        'metric_name': 'memory_usedutilization',
        'unit': '%'
    }
}

RDS_METRICS = {
    'cpu': {
        'namespace': 'acs_rds_dashboard',
        'metric_name': 'CpuUsage',
        'unit': '%'
    },
    'memory': {
        'namespace': 'acs_rds_dashboard',
        'metric_name': 'MemoryUsage',
        'unit': '%'
    }
}

def get_sprint_input():
    """Get sprint number from user input or return None for automatic calculation"""
    while True:
        sprint_input = input("\nEnter sprint number (e.g., 15) or press Enter for automatic calculation: ").strip()
        
        if not sprint_input:  # Empty input
            return None
            
        try:
            sprint_number = int(sprint_input)
            if sprint_number < 1:
                print("Sprint number must be positive.")
                continue
            return sprint_number
        except ValueError:
            print("Please enter a valid number.")

def calculate_sprint_info(target_date, manual_sprint=None):
    """Calculate sprint number and date range based on a reference sprint"""
    if manual_sprint is not None:
        # Calculate dates based on manual sprint number
        sprints_diff = manual_sprint - 15  # Difference from reference sprint
        sprint_start = SPRINT_15_START + (sprints_diff * SPRINT_DURATION)
        sprint_start = sprint_start.replace(hour=0, minute=0, second=0, microsecond=0)
        
        sprint_end = sprint_start + SPRINT_DURATION - timedelta(days=1)
        sprint_end = sprint_end.replace(hour=23, minute=59, second=59, microsecond=999999)
        
        return {
            'sprint_number': manual_sprint,
            'start_date': sprint_start,
            'end_date': sprint_end
        }
    
    # Original automatic calculation logic
    if isinstance(target_date, str):
        target_date = datetime.strptime(target_date, '%Y-%m-%d %H:%M:%S').date()
    elif isinstance(target_date, datetime):
        target_date = target_date.date()
    
    target_datetime = datetime.combine(target_date, datetime.min.time().replace(hour=12))
    days_diff = (target_datetime - SPRINT_15_START).days
    sprints_diff = days_diff // SPRINT_DURATION.days
    
    sprint_start = SPRINT_15_START + (sprints_diff * SPRINT_DURATION)
    sprint_start = sprint_start.replace(hour=0, minute=0, second=0, microsecond=0)
    
    sprint_end = sprint_start + SPRINT_DURATION - timedelta(days=1)
    sprint_end = sprint_end.replace(hour=23, minute=59, second=59, microsecond=999999)
    
    current_sprint = 15 + sprints_diff
    
    return {
        'sprint_number': current_sprint,
        'start_date': sprint_start,
        'end_date': sprint_end
    }

def get_cloud_monitor_data(client, namespace, metric_name, instance_id, start_time=None, end_time=None):
    """Fetch monitoring data from Alibaba Cloud Monitor"""
    if start_time is None or end_time is None:
        end_time = datetime.now()
        start_time = end_time - timedelta(days=14)
    
    print(f"\nRequesting data for {metric_name} (Instance: {instance_id}):")
    print(f"Start time: {start_time}")
    print(f"End time: {end_time}")
    
    # Split the time range into 3-day chunks to ensure we get all data
    chunk_size = timedelta(days=3)
    current_start = start_time
    all_datapoints = []
    
    while current_start < end_time:
        current_end = min(current_start + chunk_size, end_time)
        
        start_timestamp = int(current_start.timestamp()) * 1000
        end_timestamp = int(current_end.timestamp()) * 1000
        
        request = DescribeMetricListRequest()
        request.set_accept_format('json')
        request.set_Namespace(namespace)
        request.set_MetricName(metric_name)
        request.set_StartTime(start_timestamp)
        request.set_EndTime(end_timestamp)
        request.set_Dimensions(f'{{"instanceId": "{instance_id}"}}')
        request.set_Period('7200')
        
        response = client.do_action_with_exception(request)
        response_data = json.loads(response)
        
        if 'Datapoints' in response_data:
            chunk_datapoints = response_data['Datapoints']
            if isinstance(chunk_datapoints, str):
                chunk_datapoints = json.loads(chunk_datapoints)
            all_datapoints.extend(chunk_datapoints)
        
        current_start = current_end
    
    return {'Datapoints': all_datapoints}

def collect_metrics(client, instances, metrics, start_time, end_time):
    """Collect metrics for multiple instances"""
    results = {}
    
    for instance in instances:
        instance_id = instance['id']
        instance_name = instance['name']
        results[instance_name] = {}
        
        for metric_type, metric_config in metrics.items():
            data = get_cloud_monitor_data(
                client,
                metric_config['namespace'],
                metric_config['metric_name'],
                instance_id,
                start_time,
                end_time
            )
            df, stats, anomalies = process_metrics(data)
            results[instance_name][metric_type] = {
                'data': df,
                'stats': stats,
                'anomalies': anomalies,
                'unit': metric_config['unit']
            }
    
    return results

def process_metrics(raw_data):    
    if 'Datapoints' not in raw_data or not raw_data['Datapoints']:
        return pd.DataFrame(), {'average': 0, 'max': 0, 'min': 0}, pd.DataFrame()
    
    if isinstance(raw_data['Datapoints'], str):
        datapoints = json.loads(raw_data['Datapoints'])
    else:
        datapoints = raw_data['Datapoints']
    
    df = pd.DataFrame(datapoints)
    df['timestamp'] = pd.to_datetime(df['timestamp'], unit='ms')
    df.set_index('timestamp', inplace=True)
    
    value_column = 'Value' if 'Value' in df.columns else 'Average'
    if value_column not in df.columns:
        return pd.DataFrame(), {'average': 0, 'max': 0, 'min': 0}, pd.DataFrame()
    
    stats = {
        'average': df[value_column].mean(),
        'max': df[value_column].max(),
        'min': df[value_column].min()
    }
    
    anomaly_threshold = 80
    anomalies = df[df[value_column] > anomaly_threshold]
    
    return df, stats, anomalies

def generate_combined_chart(data_dict, metric_type, title, filename, days=14):
    """Generate combined chart for multiple instances"""
    plt.figure(figsize=(12, 6))
    
    for instance_name, metrics in data_dict.items():
        if metric_type in metrics and not metrics[metric_type]['data'].empty:
            df = metrics[metric_type]['data']
            plt.plot(df.index, df['Average'], label=instance_name, linewidth=1)
    
    plt.title(f'{title} (Last {days} Days)')
    plt.ylabel(f'{metric_type.upper()} Utilization (%)')
    plt.xlabel('Date/Time')
    plt.grid(True)
    plt.gcf().autofmt_xdate()
    
    ax = plt.gca()
    ax.xaxis.set_major_formatter(plt.matplotlib.dates.DateFormatter('%d-%m %H:%M:%S'))
    
    plt.xticks(rotation=45)
    plt.legend()
    plt.tight_layout()
    
    # Ensure charts directory exists
    os.makedirs(CHARTS_DIR, exist_ok=True)
    filepath = os.path.join(CHARTS_DIR, filename)
    plt.savefig(filepath, dpi=150)
    plt.close()
    
    return filepath

def create_word_report(report_data, incidents=None, recommendations=None, sprint_info=None):
    """Generate Word document report with the new structure"""
    # Use provided sprint info or calculate based on current time
    if sprint_info is None:
        current_time = datetime.now()
        sprint_info = calculate_sprint_info(current_time.strftime('%Y-%m-%d %H:%M:%S'))
    
    # Format the output filename using the sprint number
    output_file = f'Sprint{sprint_info["sprint_number"]:02d}_Alibaba_Cloud_Resources_Utilization_Report.docx'
    
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Report Header with larger font size
    title = doc.add_heading('Alibaba Cloud Resources Utilization Report', 0)
    title.runs[0].font.size = Pt(24)
    
    # Add sprint subheading with specific times
    sprint_subheading = (
        f'Sprint {sprint_info["sprint_number"]} '
        f'({sprint_info["start_date"].strftime("%d %b %Y %H:%M:%S")}-'
        f'{sprint_info["end_date"].strftime("%d %b %Y %H:%M:%S")})'
    )
    doc.add_paragraph(sprint_subheading)
    
    # Report period matches sprint period
    doc.add_paragraph(
        f'Report Period: {sprint_info["start_date"].strftime("%d %b %Y %H:%M:%S")} to '
        f'{sprint_info["end_date"].strftime("%d %b %Y %H:%M:%S")}'
    )
    
    # Overall Summary
    doc.add_heading('Overall Summary', level=1)
    
    for env in ['DEV', 'UAT']:
        doc.add_heading(env, level=2)
        
        # App and Web Servers
        doc.add_heading('App and Web Servers', level=3)
        servers_data = report_data[env]['servers']
        
        for metric in ['cpu', 'memory']:
            ranges = [d[metric]['stats'] for d in servers_data.values()]
            min_val = min(r['min'] for r in ranges)
            max_val = max(r['max'] for r in ranges)
            avg_val = sum(r['average'] for r in ranges) / len(ranges)
            
            p = doc.add_paragraph(
                f'{metric.upper()} Utilization Range: {min_val:.2f}% to {max_val:.2f}% '
                f'(Average: {avg_val:.2f}%)'
            )
            p.style = doc.styles['Normal']
        
        # Database
        doc.add_heading('Database', level=3)
        db_data = report_data[env]['rds']
        for metric in ['cpu', 'memory']:
            stats = list(db_data.values())[0][metric]['stats']
            p = doc.add_paragraph(
                f'{metric.upper()} Utilization Range: {stats["min"]:.2f}% to {stats["max"]:.2f}% '
                f'(Average: {stats["average"]:.2f}%)'
            )
            p.style = doc.styles['Normal']
    
    # Incidents Section
    doc.add_heading('Incidents', level=1)
    if incidents:
        for incident in incidents:
            p = doc.add_paragraph(f'• {incident}')
            p.style = doc.styles['Normal']
    else:
        p = doc.add_paragraph('No incidents reported during this period.')
        p.style = doc.styles['Normal']
    
    # Recommendations Section
    doc.add_heading('Recommendations', level=1)
    if recommendations:
        for recommendation in recommendations:
            p = doc.add_paragraph(f'• {recommendation}')
            p.style = doc.styles['Normal']
    else:
        p = doc.add_paragraph('No specific recommendations for this period.')
        p.style = doc.styles['Normal']
    
    # Charts Sections
    for env in ['DEV', 'UAT']:
        doc.add_heading(f'{env} Dashboards', level=1)
        
        # Server Charts
        doc.add_heading('App and Web Servers', level=2)
        for metric in ['cpu', 'memory']:
            chart_path = os.path.join(CHARTS_DIR, f'{env.lower()}_servers_{metric}_chart.png')
            doc.add_picture(chart_path, width=Inches(6))
        
        # RDS Charts
        doc.add_heading('Database', level=2)
        for metric in ['cpu', 'memory']:
            chart_path = os.path.join(CHARTS_DIR, f'{env.lower()}_rds_{metric}_chart.png')
            doc.add_picture(chart_path, width=Inches(6))
    
    # Save document
    doc.save(output_file)
    print(f'Report generated: {output_file}')

def main():
    # Initialize Alibaba Cloud client
    client = AcsClient(ACCESS_KEY_ID, ACCESS_KEY_SECRET, REGION_ID)
    
    # Get manual sprint input if provided
    manual_sprint = get_sprint_input()
    
    # Calculate sprint period based on current time or manual input
    sprint_info = calculate_sprint_info(
        datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        manual_sprint
    )
    
    start_time = sprint_info['start_date']
    end_time = sprint_info['end_date']
    
    print(f"\nGenerating report for Sprint {sprint_info['sprint_number']}")
    print(f"Period: {start_time.strftime('%Y-%m-%d %H:%M:%S')} to {end_time.strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Set timezone to Hong Kong
    hk_tz = pytz.timezone('Asia/Hong_Kong')
    start_time = start_time.replace(tzinfo=hk_tz)
    end_time = end_time.replace(tzinfo=hk_tz)
    
    # Remove timezone info as Alibaba Cloud API expects UTC
    start_time = start_time.replace(tzinfo=None)
    end_time = end_time.replace(tzinfo=None)
    
    monitoring_days = 14
    report_data = {
        'start_time': start_time.strftime('%Y-%m-%d %H:%M:%S'),
        'end_time': end_time.strftime('%Y-%m-%d %H:%M:%S'),
    }
    
    # Clean up and recreate charts directory
    if os.path.exists(CHARTS_DIR):
        # If it's a file, remove it
        if os.path.isfile(CHARTS_DIR):
            os.remove(CHARTS_DIR)
        else:
            # If it's a directory, remove it and its contents
            shutil.rmtree(CHARTS_DIR)
    
    # Create fresh charts directory
    os.makedirs(CHARTS_DIR)
    
    # Collect metrics for each environment
    for env in ['DEV', 'UAT']:
        report_data[env] = {}
        
        # Collect server metrics
        servers_metrics = collect_metrics(
            client,
            INSTANCES[env]['web'] + INSTANCES[env]['app'],
            ECS_METRICS,
            start_time,
            end_time
        )
        report_data[env]['servers'] = servers_metrics
        
        # Generate combined server charts
        for metric in ['cpu', 'memory']:
            generate_combined_chart(
                servers_metrics,
                metric,
                f'{env} Servers {metric.upper()} Utilization',
                f'{env.lower()}_servers_{metric}_chart.png',
                days=monitoring_days
            )
        
        # Collect RDS metrics
        rds_metrics = collect_metrics(
            client,
            INSTANCES[env]['rds'],
            RDS_METRICS,
            start_time,
            end_time
        )
        report_data[env]['rds'] = rds_metrics
        
        # Generate RDS charts
        for metric in ['cpu', 'memory']:
            generate_combined_chart(
                rds_metrics,
                metric,
                f'{env} RDS {metric.upper()} Utilization',
                f'{env.lower()}_rds_{metric}_chart.png',
                days=monitoring_days
            )
    
    # Example incidents and recommendations (can be passed as parameters)
    incidents = None  # Add incidents here if needed
    recommendations = None  # Add recommendations here if needed
    
    # Generate Word document with sprint info
    create_word_report(report_data, incidents, recommendations, sprint_info)
    
    # Cleanup charts directory
    if os.path.exists(CHARTS_DIR):
        shutil.rmtree(CHARTS_DIR)

if __name__ == "__main__":
    main()